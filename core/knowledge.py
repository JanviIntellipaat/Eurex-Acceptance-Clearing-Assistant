from __future__ import annotations
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional, Dict, Any
import hashlib
import io
import os

from .embeddings import Embeddings
from .vector_faiss import FaissVectorStore, Chunk
from .structured import StructuredStore

try:
    import pdfplumber
except Exception:
    pdfplumber = None

try:
    import docx
except Exception:
    docx = None

@dataclass
class IndexTextResult:
    id: str
    source: str
    page: Optional[int]
    chars: int
    preview: str

class KnowledgeBase:
    def __init__(self, settings):
        self.dir = Path(settings.kb_dir); self.dir.mkdir(parents=True, exist_ok=True)
        self.embed = Embeddings(backend=settings.embed_backend, base_url=settings.base_url)
        self.vs = FaissVectorStore(self.dir / "index", Path(settings.duckdb_path), dim=settings.embed_dim)
        self.structured = StructuredStore(settings.duckdb_path)

    def _chunk_text(self, text: str, size: int = 300) -> List[str]:
        words = text.split()
        return [" ".join(words[i:i+size]) for i in range(0, len(words), size)]

    # -------- heading extractors ----------
    def _docx_headings(self, document) -> List[tuple[int,int,str]]:
        res = []
        for p in document.paragraphs:
            st = getattr(p, "style", None)
            name = getattr(st, "name", "") if st else ""
            if name.lower().startswith("heading"):
                try:
                    lvl = int("".join(ch for ch in name if ch.isdigit()))
                except Exception:
                    lvl = 1
                text = p.text.strip()
                if text:
                    res.append((1, lvl, text))
        return res

    def _pdf_headings(self, pdf) -> List[tuple[int,int,str]]:
        res = []
        if not pdfplumber: return res
        try:
            for pidx, page in enumerate(pdf.pages, start=1):
                words = page.extract_words(use_text_flow=True, extra_attrs=["size"]) or []
                if not words: continue
                sizes = [w.get("size", 0) for w in words if w.get("size", 0)]
                if not sizes: continue
                median = sorted(sizes)[len(sizes)//2]
                lines: Dict[int, List[dict]] = {}
                for w in words:
                    top = int(w.get("top", 0))
                    lines.setdefault(top, []).append(w)
                for _, ws in lines.items():
                    avg = sum(w.get("size", 0) for w in ws) / max(1, len(ws))
                    if avg >= median * 1.25:
                        txt = " ".join(w.get("text","") for w in ws).strip()
                        if txt:
                            level = 1 if avg >= median * 1.6 else 2
                            res.append((pidx, level, txt))
        except Exception:
            pass
        return res

    # -------- ingestion ----------
    def add_files(self, files, parse_tables: bool = True) -> Dict[str, Any]:
        text_results: List[IndexTextResult] = []
        tables_added = 0

        for f in files:
            name = getattr(f, "name", "upload")
            content = f.read()
            path = self.dir / name
            path.parent.mkdir(parents=True, exist_ok=True)
            with open(path, "wb") as out:
                out.write(content)

            lower = name.lower()
            size_bytes = os.path.getsize(path)

            if lower.endswith(".pdf") and pdfplumber:
                with pdfplumber.open(path) as pdf:
                    if parse_tables:
                        tables_added += self.structured.add_pdf_tables(name, pdf).tables_added
                    heads = self._pdf_headings(pdf)
                    if heads:
                        self.structured.add_headings(name, heads)
                    for i, page in enumerate(pdf.pages, start=1):
                        text = page.extract_text() or ""
                        if not text.strip():
                            continue
                        for idx, chunk in enumerate(self._chunk_text(text)):
                            uid = hashlib.sha1(f"{name}-{i}-{idx}-{len(chunk)}".encode()).hexdigest()
                            text_results.append(IndexTextResult(id=uid, source=name, page=i, chars=len(chunk), preview=chunk[:160]))
                self.structured.record_document(name, "pdf", size_bytes=size_bytes, sheet_count=0, table_count=0)

            elif lower.endswith(".docx") and docx:
                document = docx.Document(str(path))
                try:
                    res = self.structured.add_docx_tables(name, document)
                    tables_added += res.tables_added
                except Exception:
                    pass
                heads = self._docx_headings(document)
                if heads:
                    self.structured.add_headings(name, heads)
                buf = [p.text for p in document.paragraphs if p.text and p.text.strip()]
                if buf:
                    full = " ".join(buf)
                    for idx, chunk in enumerate(self._chunk_text(full)):
                        uid = hashlib.sha1(f"{name}-docx-{idx}-{len(chunk)}".encode()).hexdigest()
                        text_results.append(IndexTextResult(id=uid, source=name, page=1, chars=len(chunk), preview=chunk[:160]))
                self.structured.record_document(name, "docx", size_bytes=size_bytes, sheet_count=0, table_count=0)

            elif lower.endswith(".csv"):
                tables_added += self.structured.add_csv(name, content).tables_added
                try:
                    import pandas as pd
                    df = pd.read_csv(io.BytesIO(content), dtype=str, keep_default_na=False, low_memory=False)
                    text = df.head(50).to_csv(index=False)
                    for idx, chunk in enumerate(self._chunk_text(text, size=200)):
                        uid = hashlib.sha1(f"{name}-csv-{idx}-{len(chunk)}".encode()).hexdigest()
                        text_results.append(IndexTextResult(id=uid, source=name, page=1, chars=len(chunk), preview=chunk[:160]))
                except Exception:
                    pass

            elif lower.endswith(".xlsx"):
                tables_added += self.structured.add_xlsx(name, content).tables_added
                try:
                    import pandas as pd
                    sheets = pd.read_excel(io.BytesIO(content), dtype=str, sheet_name=None, engine="openpyxl")
                    for sname, df in list(sheets.items())[:3]:
                        text = df.head(30).to_csv(index=False)
                        for idx, chunk in enumerate(self._chunk_text(text, size=200)):
                            uid = hashlib.sha1(f"{name}-{sname}-xlsx-{idx}-{len(chunk)}".encode()).hexdigest()
                            text_results.append(IndexTextResult(id=uid, source=f"{name}#{sname}", page=1, chars=len(chunk), preview=chunk[:160]))
                except Exception:
                    pass

            elif lower.endswith(".xls"):
                tables_added += self.structured.add_xls(name, content).tables_added
                try:
                    import pandas as pd
                    sheets = pd.read_excel(io.BytesIO(content), dtype=str, sheet_name=None, engine="xlrd")
                    for sname, df in list(sheets.items())[:3]:
                        text = df.head(30).to_csv(index=False)
                        for idx, chunk in enumerate(self._chunk_text(text, size=200)):
                            uid = hashlib.sha1(f"{name}-{sname}-xls-{idx}-{len(chunk)}".encode()).hexdigest()
                            text_results.append(IndexTextResult(id=uid, source=f"{name}#{sname}", page=1, chars=len(chunk), preview=chunk[:160]))
                except Exception:
                    pass

            elif lower.endswith(".md") or lower.endswith(".txt"):
                try:
                    text = path.read_text(encoding="utf-8", errors="ignore")
                except Exception:
                    text = ""
                for idx, chunk in enumerate(self._chunk_text(text)):
                    uid = hashlib.sha1(f"{name}-raw-{idx}-{len(chunk)}".encode()).hexdigest()
                    text_results.append(IndexTextResult(id=uid, source=name, page=1, chars=len(chunk), preview=chunk[:160]))
                self.structured.record_document(name, "text", size_bytes=size_bytes, sheet_count=0, table_count=0)

            else:
                try:
                    text = path.read_text(encoding="utf-8", errors="ignore")
                    if text.strip():
                        for idx, chunk in enumerate(self._chunk_text(text)):
                            uid = hashlib.sha1(f"{name}-raw-{idx}-{len(chunk)}".encode()).hexdigest()
                            text_results.append(IndexTextResult(id=uid, source=name, page=1, chars=len(chunk), preview=chunk[:160]))
                except Exception:
                    pass
                self.structured.record_document(name, "other", size_bytes=size_bytes, sheet_count=0, table_count=0)

        if text_results:
            docs = [t.preview for t in text_results]
            vecs = self.embed.embed(docs)
            ids = [t.id for t in text_results]
            metadatas = [{"source": t.source, "page": t.page, "text": t.preview} for t in text_results]
            self.vs.add(ids, vecs, metadatas)

        return {"text_chunks": [t.__dict__ for t in text_results], "tables_added": tables_added}

    # retrieval
    def search_text(self, query: str, k: int = 5) -> List[Chunk]:
        if not query.strip():
            return []
        vec = self.embed.embed([query])
        return self.vs.query(vec, k=k)

    def search_structured_context(self, query: str, top_k: int = 5) -> str:
        return self.structured.search_context(query, top_k=top_k)

    # admin
    def purge_index(self):
        self.vs.purge()

    def purge_structured(self):
        self.structured.purge()
